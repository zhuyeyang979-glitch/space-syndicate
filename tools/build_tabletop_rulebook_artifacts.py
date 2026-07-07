from __future__ import annotations

import html
import math
import re
import textwrap
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "tabletop_rulebook.md"
ASSET_DIR = ROOT / "tmp" / "rulebook_assets"
DOCX_OUT = ROOT / "output" / "documents" / "space_syndicate_tabletop_rulebook_v0.2.docx"
PDF_OUT = ROOT / "output" / "pdf" / "space_syndicate_tabletop_rulebook_v0.2.pdf"

FONT_BODY = Path("C:/Windows/Fonts/Deng.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/Dengb.ttf")
FONT_HEI = Path("C:/Windows/Fonts/simhei.ttf")

NAVY = "0B2545"
BLUE = "2E74B5"
CYAN = "41C7D4"
GOLD = "F3C969"
ORANGE = "D9893D"
RED = "B64B4B"
GREEN = "4A9A68"
GRAY = "6B7280"
LIGHT = "F4F7FB"
INK = "172033"
LINE = "C6D1DF"


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)


def load_font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def wrap_cjk(text: str, width: int) -> list[str]:
    chunks: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            line += ch
            if len(line) >= width:
                chunks.append(line)
                line = ""
        if line:
            chunks.append(line)
    return chunks or [""]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = INK,
    max_chars: int = 10,
    line_gap: int = 6,
) -> None:
    lines = wrap_cjk(text, max_chars)
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + max(0, len(lines) - 1) * line_gap
    x1, y1, x2, y2 = box
    y = y1 + (y2 - y1 - total_h) / 2
    for idx, line in enumerate(lines):
        x = x1 + (x2 - x1 - widths[idx]) / 2
        draw.text((x, y), line, font=font, fill=hex_rgb(fill))
        y += line_heights[idx] + line_gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str = LINE,
    width: int = 3,
    radius: int = 24,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=hex_rgb(fill), outline=hex_rgb(outline), width=width)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = BLUE, width: int = 5) -> None:
    draw.line([start, end], fill=hex_rgb(fill), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    points = [
        end,
        (int(end[0] - size * math.cos(angle - math.pi / 6)), int(end[1] - size * math.sin(angle - math.pi / 6))),
        (int(end[0] - size * math.cos(angle + math.pi / 6)), int(end[1] - size * math.sin(angle + math.pi / 6))),
    ]
    draw.polygon(points, fill=hex_rgb(fill))


def base_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 850), hex_rgb("FFFFFF"))
    draw = ImageDraw.Draw(img)
    title_font = load_font(FONT_BOLD, 48)
    sub_font = load_font(FONT_BODY, 24)
    draw.rounded_rectangle((34, 34, 1566, 816), radius=36, fill=hex_rgb("FFFFFF"), outline=hex_rgb("D7DEE8"), width=3)
    draw.rectangle((34, 34, 1566, 145), fill=hex_rgb("EEF6FF"))
    draw.text((72, 58), title, font=title_font, fill=hex_rgb(NAVY))
    draw.text((72, 112), subtitle, font=sub_font, fill=hex_rgb(GRAY))
    return img, draw


def diagram_information() -> Path:
    img, draw = base_canvas("信息边界", "玩家靠公开线索推理，不直接看到对手的私有资源")
    font = load_font(FONT_BOLD, 30)
    small = load_font(FONT_BODY, 24)
    boxes = [
        ((110, 220, 500, 640), "公开桌面", ["星球地图", "城市/破坏", "牌轨/竞价", "价格/天气", "已揭示归属"], "EAF4FF"),
        ((605, 220, 995, 640), "你的私有区", ["准确现金", "手牌", "弃牌选择", "城市归属", "私有标注"], "FFF7E6"),
        ((1100, 220, 1490, 640), "对手隐藏区", ["准确现金不可见", "手牌不可见", "弃牌不可见", "计划不可见", "只能推理"], "F3F4F6"),
    ]
    for box, title, items, fill in boxes:
        rounded_box(draw, box, fill, "CAD6E3", width=4)
        draw_centered_text(draw, (box[0] + 20, box[1] + 24, box[2] - 20, box[1] + 92), title, font, NAVY, 8)
        y = box[1] + 124
        for item in items:
            draw.text((box[0] + 58, y), f"- {item}", font=small, fill=hex_rgb(INK))
            y += 56
    arrow(draw, (500, 430), (605, 430), ORANGE)
    arrow(draw, (1100, 430), (995, 430), ORANGE)
    draw.text((660, 690), "下注、目标、损伤、商品条件会变成可读线索", font=small, fill=hex_rgb(BLUE))
    path = ASSET_DIR / "diagram_01_information.png"
    img.save(path)
    return path


def diagram_core_loop() -> Path:
    img, draw = base_canvas("一局核心循环", "实时进行，但玩家可以按桌游动作节奏读桌面")
    font = load_font(FONT_BOLD, 26)
    steps = [
        ("首召怪兽", "进入购牌网络"),
        ("城市化牌", "新增生产/需求/通商"),
        ("购买卡牌", "从区域牌架扩张路线"),
        ("公开出牌", "结算、竞价、留下线索"),
        ("经济变化", "GDP/商品/商路转成现金"),
        ("怪兽/军队压力", "破坏、保护、下注、推理"),
    ]
    centers = [(290, 300), (650, 300), (1010, 300), (1190, 575), (650, 575), (290, 575)]
    for idx, ((title, sub), center) in enumerate(zip(steps, centers)):
        x, y = center
        fill = ["EAF4FF", "EAFBF0", "FFF7E6", "FCECEB", "F1F0FF", "EEF6FF"][idx]
        rounded_box(draw, (x - 140, y - 80, x + 140, y + 80), fill, "B8C7D9", width=4)
        draw_centered_text(draw, (x - 128, y - 62, x + 128, y - 8), title, font, NAVY, 8)
        draw_centered_text(draw, (x - 118, y + 4, x + 118, y + 58), sub, load_font(FONT_BODY, 20), GRAY, 9)
    for a, b in zip(centers, centers[1:] + centers[:1]):
        start = (int(a[0] + (b[0] - a[0]) * 0.28), int(a[1] + (b[1] - a[1]) * 0.28))
        end = (int(a[0] + (b[0] - a[0]) * 0.72), int(a[1] + (b[1] - a[1]) * 0.72))
        arrow(draw, start, end, BLUE)
    draw.text((585, 722), "目标：终局倒计时结束时现金最多", font=load_font(FONT_BOLD, 28), fill=hex_rgb(ORANGE))
    path = ASSET_DIR / "diagram_02_core_loop.png"
    img.save(path)
    return path


def diagram_gdp() -> Path:
    img, draw = base_canvas("城市化份额如何变成现金", "城市化牌新增商品项目，收益按隐藏份额分配")
    font = load_font(FONT_BOLD, 25)
    small = load_font(FONT_BODY, 21)
    left_nodes = [
        ((90, 215, 355, 310), "生产项目", "某商品产出"),
        ((90, 345, 355, 440), "需求项目", "消费/订单"),
        ((90, 475, 355, 570), "通商项目", "速度/商路"),
        ((90, 605, 355, 700), "合约项目", "按商品连接"),
    ]
    right_nodes = [
        ((1215, 250, 1490, 350), "份额饼图", "你 vs 对手合计"),
        ((1215, 405, 1490, 505), "商品控制者", "决定相关合约"),
        ((1215, 560, 1490, 660), "损伤/市场", "影响项目 GDP"),
    ]
    for box, title, sub in left_nodes + right_nodes:
        rounded_box(draw, box, "F8FAFC", "CBD5E1", width=3)
        draw_centered_text(draw, (box[0] + 12, box[1] + 8, box[2] - 12, box[1] + 48), title, font, NAVY, 6)
        draw_centered_text(draw, (box[0] + 12, box[1] + 48, box[2] - 12, box[3] - 8), sub, small, GRAY, 8)
    rounded_box(draw, (555, 315, 1045, 500), "EAFBF0", "80B88A", width=5)
    draw_centered_text(draw, (585, 334, 1015, 400), "城市化项目 GDP", load_font(FONT_BOLD, 38), GREEN, 10)
    draw_centered_text(draw, (585, 405, 1015, 480), "生产/需求/通商各自计算", load_font(FONT_BODY, 27), NAVY, 11)
    rounded_box(draw, (560, 600, 1040, 725), "FFF7E6", "E4B454", width=4)
    draw_centered_text(draw, (580, 615, 1020, 675), "按城市化份额分钱", load_font(FONT_BOLD, 32), ORANGE, 10)
    draw_centered_text(draw, (580, 680, 1020, 713), "金融协议也按相关份额结算", load_font(FONT_BODY, 21), NAVY, 14)
    for box, _, _ in left_nodes:
        arrow(draw, (box[2], (box[1] + box[3]) // 2), (555, 410), GREEN)
    for box, _, _ in right_nodes:
        arrow(draw, (1045, 410), (box[0], (box[1] + box[3]) // 2), RED)
    arrow(draw, (800, 500), (800, 600), ORANGE)
    path = ASSET_DIR / "diagram_03_gdp.png"
    img.save(path)
    return path


def diagram_card_track() -> Path:
    img, draw = base_canvas("公开牌轨与竞价", "所有人看得到牌和结果，但默认不知道是谁打的")
    font = load_font(FONT_BOLD, 25)
    small = load_font(FONT_BODY, 20)
    boxes = [
        ((80, 310, 300, 470), "手牌", "选择目标\n提交即承诺", "FFF7E6"),
        ((380, 310, 600, 470), "公开展示", "卡牌/目标/条件", "EAF4FF"),
        ((680, 310, 900, 470), "竞价排序", "金额公开\n身份隐藏", "F1F0FF"),
        ((980, 310, 1200, 470), "结算", "效果落到地图\n经济变化", "EAFBF0"),
        ((1280, 310, 1500, 470), "推理", "猜归属\n赢钱/贴标签", "FCECEB"),
    ]
    for box, title, sub, fill in boxes:
        rounded_box(draw, box, fill, "B8C7D9", width=4)
        draw_centered_text(draw, (box[0] + 10, box[1] + 20, box[2] - 10, box[1] + 72), title, font, NAVY, 7)
        draw_centered_text(draw, (box[0] + 10, box[1] + 78, box[2] - 10, box[3] - 18), sub, small, GRAY, 7)
    for box_a, box_b in zip(boxes, boxes[1:]):
        arrow(draw, (box_a[0][2] + 18, 390), (box_b[0][0] - 18, 390), BLUE)
    draw.text((330, 610), "同一短窗内多张牌才进入竞价；单张牌直接展示并结算。", font=load_font(FONT_BODY, 27), fill=hex_rgb(NAVY))
    path = ASSET_DIR / "diagram_04_card_track.png"
    img.save(path)
    return path


def diagram_monster() -> Path:
    img, draw = base_canvas("怪兽压力系统", "怪兽自动行动，玩家只能用卡牌制造一次性影响")
    font = load_font(FONT_BOLD, 25)
    small = load_font(FONT_BODY, 20)
    center = (630, 340, 980, 535)
    rounded_box(draw, center, "FCECEB", "D16A6A", width=5)
    draw_centered_text(draw, (660, 365, 950, 435), "自动怪兽", load_font(FONT_BOLD, 42), RED, 6)
    draw_centered_text(draw, (660, 438, 950, 510), "概率行动表 + 生态偏好", small, NAVY, 10)
    inputs = [
        ((90, 220, 355, 315), "商品偏好", "资源越吻合越想去"),
        ((90, 385, 355, 480), "城市/GDP", "高价值城市更诱人"),
        ((90, 550, 355, 645), "天气/地形", "海栖、飞行、陆行差异"),
    ]
    outputs = [
        ((1215, 220, 1490, 315), "破坏", "区域/城市/商路受损"),
        ((1215, 385, 1490, 480), "线索", "受伤会牵动归属现金"),
        ((1215, 550, 1490, 645), "赌局", "怪兽交战触发下注"),
    ]
    for box, title, sub in inputs + outputs:
        rounded_box(draw, box, "F8FAFC", "CBD5E1", width=3)
        draw_centered_text(draw, (box[0] + 10, box[1] + 10, box[2] - 10, box[1] + 48), title, font, NAVY, 7)
        draw_centered_text(draw, (box[0] + 10, box[1] + 48, box[2] - 10, box[3] - 8), sub, small, GRAY, 8)
    for box, _, _ in inputs:
        arrow(draw, (box[2], (box[1] + box[3]) // 2), (center[0], 438), ORANGE)
    for box, _, _ in outputs:
        arrow(draw, (center[2], 438), (box[0], (box[1] + box[3]) // 2), RED)
    rounded_box(draw, (560, 645, 1055, 735), "FFF7E6", "E4B454", width=3)
    draw_centered_text(draw, (580, 657, 1035, 718), "绑定技能/诱导牌：只影响一次具体行动", load_font(FONT_BOLD, 27), ORANGE, 14)
    arrow(draw, (808, 645), (808, 535), ORANGE)
    path = ASSET_DIR / "diagram_05_monster_pressure.png"
    img.save(path)
    return path


def diagram_wager() -> Path:
    img, draw = base_canvas("怪兽赌局", "怪兽战斗前冻结时间，全员按现金百分比公开下注")
    font = load_font(FONT_BOLD, 24)
    small = load_font(FONT_BODY, 19)
    steps = [
        ("怪兽相遇", "准备战斗"),
        ("时间冻结", "最多 30 秒"),
        ("基础比例", "现金的 5%-10%"),
        ("公开加注", "每次 +1%"),
        ("选择阵营", "押造成最多伤害者"),
        ("赢家分池", "逆风翻盘机会"),
    ]
    x = 85
    for idx, (title, sub) in enumerate(steps):
        box = (x, 345, x + 205, 505)
        fill = ["FCECEB", "F1F0FF", "FFF7E6", "EAF4FF", "F8FAFC", "EAFBF0"][idx]
        rounded_box(draw, box, fill, "B8C7D9", width=4)
        draw_centered_text(draw, (box[0] + 8, box[1] + 22, box[2] - 8, box[1] + 76), title, font, NAVY, 6)
        draw_centered_text(draw, (box[0] + 8, box[1] + 84, box[2] - 8, box[3] - 18), sub, small, GRAY, 7)
        if idx < len(steps) - 1:
            arrow(draw, (x + 220, 425), (x + 270, 425), BLUE)
        x += 250
    draw.text((180, 640), "下注金额和玩家身份公开，所以赌局也是推理线索；但胜负按战斗结果结算。", font=load_font(FONT_BODY, 27), fill=hex_rgb(NAVY))
    path = ASSET_DIR / "diagram_06_wager.png"
    img.save(path)
    return path


def build_diagrams() -> dict[str, tuple[Path, str]]:
    return {
        "## 3. 玩家人数与信息边界": (diagram_information(), "图 1：公开信息、你的私有信息与对手隐藏信息的边界。"),
        "## 6. 一局游戏的基本流程": (diagram_core_loop(), "图 2：一局游戏的核心循环。"),
        "## 7. 城市、GDP 与赚钱": (diagram_gdp(), "图 3：城市化项目、份额饼图与 GDP 分配。"),
        "## 10. 出牌、公开牌轨与竞价": (diagram_card_track(), "图 4：卡牌从手牌到公开牌轨、竞价、结算和推理的路径。"),
        "## 12. 怪兽规则": (diagram_monster(), "图 5：怪兽自动目标、玩家一次性影响和公开线索。"),
        "## 13. 怪兽赌局": (diagram_wager(), "图 6：怪兽赌局的冻结、百分比下注与结算。"),
    }


def split_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def parse_markdown(path: Path) -> list[dict]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append({"type": "heading", "level": level, "text": line[level:].strip(), "raw": line})
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[i + 1]):
            rows = [split_table_row(line)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        if line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append({"type": "numbers", "items": items})
            continue
        paras = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("- ") and not re.match(r"^\d+\.\s+", lines[i]) and not lines[i].startswith("|"):
            paras.append(lines[i].strip())
            i += 1
        blocks.append({"type": "para", "text": " ".join(paras)})
    return blocks


def set_run_font(run, name: str = "DengXian", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style(style, name: str, size: float, color: str, bold: bool = False, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "DengXian", 9.5, color, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_callout_docx(doc: Document, text: str, fill: str = "F4F7FB") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.25])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "DengXian", 10.5, NAVY, True)


def add_diagram_docx(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles["Caption"]


def add_table_docx(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    norm = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(norm), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if cols == 2:
        widths = [1.55, 4.7]
    elif cols == 3:
        widths = [1.8, 2.25, 2.2]
    else:
        widths = [6.25 / cols] * cols
    set_table_widths(table, widths)
    for r_idx, row in enumerate(norm):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                shade_cell(cell, "E8EEF5")
            set_cell_text(cell, text, bold=(r_idx == 0), color=NAVY if r_idx == 0 else INK)
    doc.add_paragraph()


def build_docx(blocks: list[dict], diagrams: dict[str, tuple[Path, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], "DengXian", 10.5, INK, False, 0, 6, 1.25)
    set_style(doc.styles["Heading 1"], "DengXian", 16, BLUE, True, 18, 10, 1.15)
    set_style(doc.styles["Heading 2"], "DengXian", 13, BLUE, True, 14, 7, 1.15)
    set_style(doc.styles["Heading 3"], "DengXian", 12, NAVY, True, 10, 5, 1.15)
    set_style(doc.styles["Caption"], "DengXian", 9, GRAY, False, 0, 8, 1.15)
    set_style(doc.styles["List Bullet"], "DengXian", 10.5, INK, False, 0, 4, 1.25)
    set_style(doc.styles["List Number"], "DengXian", 10.5, INK, False, 0, 4, 1.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("太空辛迪加 | 试玩规则书 v0.2")
    set_run_font(run, "DengXian", 9, GRAY, False)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("玩家规则书 - 当前 Godot 4.7 原型")
    set_run_font(f_run, "DengXian", 9, GRAY, False)

    # Editorial-cover inspired first page.
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("试玩规则书")
    set_run_font(r, "DengXian", 13, ORANGE, True)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《太空辛迪加》")
    set_run_font(r, "DengXian", 30, NAVY, True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实时 PVE 隐藏信息商业怪兽牌局")
    set_run_font(r, "DengXian", 15, GRAY, False)
    p.paragraph_format.space_after = Pt(28)
    add_callout_docx(
        doc,
        "城市化份额让你赚钱；怪兽和军队让别人少赚钱；公开线索让你猜出谁在背后操作。最后，钱最多的人赢。",
        "FFF7E6",
    )
    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [1.7, 4.55])
    meta_rows = [
        ("版本", "v0.2 - 真人试玩规则书"),
        ("适用", "当前 Godot 4.7 实时 PVE roguelike 原型"),
        ("玩家", "3-8 席；通常 1 名真人玩家对 2-7 名 AI"),
        ("目标", "达到现金目标后进入终局倒计时；结束时现金最多获胜"),
    ]
    for idx, (k, v) in enumerate(meta_rows):
        shade_cell(meta.cell(idx, 0), "E8EEF5")
        set_cell_text(meta.cell(idx, 0), k, True, NAVY)
        set_cell_text(meta.cell(idx, 1), v, False, INK)

    doc.add_page_break()
    doc.add_heading("阅读路线", level=1)
    add_callout_docx(doc, "第一次试玩：先读 1-6 章理解桌面，再读 7-13 章理解赚钱、出牌和怪兽。其余章节可在遇到对应系统时查阅。", "EAF4FF")
    toc_items = [
        "1-3 章：游戏概念、目标、信息边界",
        "4-6 章：组件、开局、一局基本流程",
        "7-11 章：GDP、商品、购牌、出牌和卡牌类别",
        "12-18 章：怪兽、赌局、军队、合约、情报、天气和金融",
        "19-22 章：结束条件、第一局建议、关键词速查",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    for block in blocks[1:]:
        if block["type"] == "heading":
            level = block["level"]
            text = block["text"]
            if level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
            if block.get("raw") in diagrams:
                add_diagram_docx(doc, *diagrams[block["raw"]])
        elif block["type"] == "para":
            doc.add_paragraph(block["text"])
        elif block["type"] == "bullets":
            for item in block["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif block["type"] == "numbers":
            for item in block["items"]:
                doc.add_paragraph(item, style="List Number")
        elif block["type"] == "table":
            add_table_docx(doc, block["rows"])

    doc.save(DOCX_OUT)


def register_pdf_fonts() -> tuple[str, str]:
    pdfmetrics.registerFont(TTFont("DengXian", str(FONT_BODY)))
    pdfmetrics.registerFont(TTFont("DengXianBold", str(FONT_BOLD)))
    return "DengXian", "DengXianBold"


def ptext(text: str) -> str:
    return html.escape(text).replace("《", "《").replace("》", "》")


def add_pdf_table(story: list, rows: list[list[str]], styles: dict) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    norm = [row + [""] * (cols - len(row)) for row in rows]
    if cols == 2:
        widths = [1.55 * inch, 4.65 * inch]
    elif cols == 3:
        widths = [1.75 * inch, 2.25 * inch, 2.2 * inch]
    else:
        widths = [6.2 * inch / cols] * cols
    data = []
    for r_idx, row in enumerate(norm):
        data.append([Paragraph(ptext(cell), styles["table_header" if r_idx == 0 else "table_cell"]) for cell in row])
    table = Table(data, colWidths=widths, hAlign="CENTER", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C6D1DF")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def build_pdf(blocks: list[dict], diagrams: dict[str, tuple[Path, str]]) -> None:
    body_font, bold_font = register_pdf_fonts()
    stylesheet = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "BodyCN",
            parent=stylesheet["BodyText"],
            fontName=body_font,
            fontSize=10.3,
            leading=15.2,
            textColor=colors.HexColor("#172033"),
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "H1CN",
            fontName=bold_font,
            fontSize=16,
            leading=21,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "H2CN",
            fontName=bold_font,
            fontSize=13,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "caption": ParagraphStyle(
            "CaptionCN",
            fontName=body_font,
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
            spaceAfter=10,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "BulletCN",
            fontName=body_font,
            fontSize=10.1,
            leading=14.5,
            leftIndent=14,
            bulletIndent=2,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "table_cell": ParagraphStyle(
            "TableCellCN",
            fontName=body_font,
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor("#172033"),
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "TableHeaderCN",
            fontName=bold_font,
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#0B2545"),
            wordWrap="CJK",
        ),
        "callout": ParagraphStyle(
            "CalloutCN",
            fontName=bold_font,
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_LEFT,
            wordWrap="CJK",
        ),
        "cover_title": ParagraphStyle(
            "CoverTitleCN",
            fontName=bold_font,
            fontSize=30,
            leading=38,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitleCN",
            fontName=body_font,
            fontSize=14,
            leading=20,
            textColor=colors.HexColor("#6B7280"),
            alignment=TA_CENTER,
            spaceAfter=20,
            wordWrap="CJK",
        ),
    }

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=1 * inch,
        leftMargin=1 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
        title="太空辛迪加试玩规则书 v0.2",
        author="Space Syndicate",
    )

    def on_page(canvas, document):
        canvas.saveState()
        canvas.setFont("DengXian", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawString(1 * inch, 0.45 * inch, "太空辛迪加 | 试玩规则书 v0.2")
        canvas.drawRightString(7.5 * inch, 0.45 * inch, f"Page {document.page}")
        canvas.restoreState()

    story: list = []
    story.append(Spacer(1, 1.15 * inch))
    story.append(Paragraph("试玩规则书", ParagraphStyle("Kicker", fontName=bold_font, fontSize=13, leading=18, textColor=colors.HexColor("#D9893D"), alignment=TA_CENTER, spaceAfter=14)))
    story.append(Paragraph("《太空辛迪加》", styles["cover_title"]))
    story.append(Paragraph("实时 PVE 隐藏信息商业怪兽牌局", styles["cover_subtitle"]))
    callout_data = [[Paragraph("城市化份额让你赚钱；怪兽和军队让别人少赚钱；公开线索让你猜出谁在背后操作。最后，钱最多的人赢。", styles["callout"])]]
    callout = Table(callout_data, colWidths=[6.1 * inch], hAlign="CENTER")
    callout.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF7E6")), ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E4B454")), ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12), ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10)]))
    story.append(callout)
    story.append(Spacer(1, 0.3 * inch))
    add_pdf_table(
        story,
        [
            ["版本", "v0.2 - 真人试玩规则书"],
            ["适用", "当前 Godot 4.7 实时 PVE roguelike 原型"],
            ["玩家", "3-8 席；通常 1 名真人玩家对 2-7 名 AI"],
            ["目标", "达到现金目标后进入终局倒计时；结束时现金最多获胜"],
        ],
        styles,
    )
    story.append(PageBreak())
    story.append(Paragraph("阅读路线", styles["h1"]))
    story.append(Paragraph("第一次试玩：先读 1-6 章理解桌面，再读 7-13 章理解赚钱、出牌和怪兽。其余章节可在遇到对应系统时查阅。", styles["body"]))
    for item in [
        "1-3 章：游戏概念、目标、信息边界",
        "4-6 章：组件、开局、一局基本流程",
        "7-11 章：GDP、商品、购牌、出牌和卡牌类别",
        "12-18 章：怪兽、赌局、军队、合约、情报、天气和金融",
        "19-22 章：结束条件、第一局建议、关键词速查",
    ]:
        story.append(Paragraph(ptext(item), styles["bullet"], bulletText="•"))
    story.append(Spacer(1, 10))

    for block in blocks[1:]:
        if block["type"] == "heading":
            if block["level"] == 2:
                story.append(Paragraph(ptext(block["text"]), styles["h1"]))
            else:
                story.append(Paragraph(ptext(block["text"]), styles["h2"]))
            if block.get("raw") in diagrams:
                image_path, caption = diagrams[block["raw"]]
                story.append(
                    KeepTogether(
                        [
                            RLImage(str(image_path), width=6.25 * inch, height=3.32 * inch),
                            Paragraph(ptext(caption), styles["caption"]),
                        ]
                    )
                )
        elif block["type"] == "para":
            story.append(Paragraph(ptext(block["text"]), styles["body"]))
        elif block["type"] == "bullets":
            items = [ListItem(Paragraph(ptext(item), styles["bullet"]), bulletColor=colors.HexColor("#2E74B5")) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=12, bulletFontName=body_font))
        elif block["type"] == "numbers":
            items = [ListItem(Paragraph(ptext(item), styles["bullet"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="1", leftIndent=12, bulletFontName=body_font))
        elif block["type"] == "table":
            add_pdf_table(story, block["rows"], styles)

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def structural_audit() -> None:
    for path in [DOCX_OUT, PDF_OUT]:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Output missing or too small: {path}")
    with zipfile.ZipFile(DOCX_OUT) as zf:
        names = set(zf.namelist())
        if "word/document.xml" not in names:
            raise RuntimeError("DOCX missing document.xml")
        media = [name for name in names if name.startswith("word/media/")]
        if len(media) < 6:
            raise RuntimeError(f"Expected at least 6 embedded diagrams, found {len(media)}")
        text = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        forbidden = ["守护者", "D6", "3x3", "充能", "开发历史", "压力桶", "真实手牌", "真实现金"]
        for token in forbidden:
            if token in text:
                raise RuntimeError(f"Forbidden player-facing token in DOCX: {token}")


def main() -> None:
    ensure_dirs()
    diagrams = build_diagrams()
    blocks = parse_markdown(SOURCE_MD)
    build_docx(blocks, diagrams)
    build_pdf(blocks, diagrams)
    structural_audit()
    print(f"DOCX: {DOCX_OUT}")
    print(f"PDF:  {PDF_OUT}")
    print(f"Diagrams: {len(diagrams)}")


if __name__ == "__main__":
    main()
